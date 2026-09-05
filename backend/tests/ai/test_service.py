import asyncio
import json
from pathlib import Path
import httpx
import pytest
from app.ai.service import AIService, detect_intent
from app.ai.providers import CompatibleProvider
from app.ai.ui import validate_ui
from app.ai.extraction import extract_text


def run(coroutine):
    return asyncio.run(coroutine)


def output(**values):
    return {"intent": "EXPLORATION", "text": "可先整理目前理解。", "ui": [{"schema_version": "1.0", "type": "text", "text": "可先整理目前理解。"}], "hypothesis": None, "observation": None, "reflection": None, **values}


def service(handler, charge=None):
    settings = {"AI_PROVIDER": "compatible", "CHAT_BASE_URL": "https://chat.test/v1", "CHAT_MODEL": "test", "CHAT_API_KEY": "test-key", "EMBEDDING_BASE_URL": "https://embed.test/v1", "EMBEDDING_MODEL": "embedding", "EMBEDDING_API_KEY": "other-key", "EMBEDDING_DIMENSION": 3}
    async def reserve(kind):
        if charge is not None:
            await charge(kind)
    result = AIService(settings, reserve)
    result.provider = CompatibleProvider(settings, reserve, httpx.MockTransport(handler))
    return result


def completion(payload):
    return httpx.Response(200, json={"choices": [{"message": {"content": json.dumps(payload)}}]})


@pytest.mark.parametrize("case", json.loads(Path(__file__).with_name("eval_cases.json").read_text(encoding="utf-8")))
def test_synthetic_eval(case):
    ai = AIService()
    request = {"message": case["message"], "context": case.get("context", []), "consecutive_questions": case.get("consecutive_questions", 0)}
    result = run(ai.respond(request))
    assert result["intent"] == case["intent"]
    assert result["consecutive_questions"] <= 3
    assert result["metadata"]["fake"]
    assert "insight" not in result
    for ui in result["ui"]:
        validate_ui(ui)
    if case.get("no_question"):
        assert not any(x["type"] == "question" for x in result["ui"])
    if not request["context"]:
        assert result["hypothesis"] is None


def test_unsafe_citations_repair_charged_each_call():
    calls = []
    async def charge(kind): calls.append(kind)
    bad = output(hypothesis={"statement": "guess", "evidence_refs": ["foreign"], "counter_evidence_refs": [], "uncertainty": "unknown"})
    ai = service(lambda request: completion(bad), charge)
    result = run(ai.respond({"message": "探索", "context": [{"id": "owned"}]}))
    assert calls == ["chat", "chat"]
    assert result["hypothesis"] is None and result["metadata"]["fallback"]


def test_repair_once_then_valid():
    count = []
    def handler(request):
        count.append(request)
        return completion({"invalid": True} if len(count) == 1 else output())
    result = run(service(handler).respond({"message": "探索"}))
    assert len(count) == 2 and not result["metadata"]["fallback"]


def test_quota_failure_propagates_before_network():
    calls = []
    async def charge(kind): raise RuntimeError("quota exhausted")
    ai = service(lambda request: calls.append(request), charge)
    with pytest.raises(RuntimeError, match="quota"):
        run(ai.respond({"message": "探索"}))
    assert calls == []


def test_model_failure_safe_fallback():
    result = run(service(lambda request: httpx.Response(503)).respond({"message": "探索"}))
    assert result["metadata"]["fallback"] and result["hypothesis"] is None


def test_malicious_ui_rejected():
    with pytest.raises(ValueError):
        validate_ui({"schema_version": "1.0", "type": "text", "text": "<script>alert(1)</script>"})
    with pytest.raises(ValueError):
        validate_ui({"schema_version": "1.0", "type": "text", "text": "hi", "onclick": "evil"})


def test_question_limit_even_model_ignores_instruction():
    question = output(text="再說說？", ui=[{"schema_version": "1.0", "type": "question", "text": "再說說？", "purpose": "clarify", "target_uncertainty": "experience"}])
    result = run(service(lambda request: completion(question)).respond({"message": "探索", "consecutive_questions": 3}))
    assert result["consecutive_questions"] == 0 and result["ui"][0]["type"] == "choice_cards"


def test_fatigue_does_not_call_model():
    def handler(request): raise AssertionError("must not call")
    result = run(service(handler).respond({"message": "我累了，不要再問"}))
    assert result["consecutive_questions"] == 0


def test_rerank_fallback_interleave_and_dedupe():
    ai = service(lambda request: completion({"ids": ["unowned"]}))
    semantic = [{"id": "a"}, {"id": "b"}]
    temporal = [{"id": "b"}, {"id": "c"}]
    assert run(ai.rerank("q", semantic, temporal)) == [{"id": "a"}, {"id": "b"}, {"id": "c"}]


def test_rerank_exact_order_and_charge():
    calls = []
    async def charge(kind): calls.append(kind)
    ai = service(lambda request: completion({"ids": ["b", "a"]}), charge)
    assert [x["id"] for x in run(ai.rerank("q", [{"id": "a"}], [{"id": "b"}]))] == ["b", "a"]
    assert calls == ["rerank"]


def test_embedding_endpoint_dimension_and_charge():
    calls = []
    async def charge(kind): calls.append(kind)
    def handler(request):
        assert request.url.host == "embed.test"
        assert request.headers["authorization"] == "Bearer other-key"
        return httpx.Response(200, json={"data": [{"index": 0, "embedding": [1, 0, 0]}]})
    ai = service(handler, charge)
    assert run(ai.embed(["hello"])) == [[1, 0, 0]]
    assert calls == ["embedding"]
    ai.dimension = 4
    with pytest.raises(ValueError, match="dimension"):
        run(ai.embed(["hello"]))


def test_fake_embedding_stable_and_finite():
    ai = AIService({"EMBEDDING_DIMENSION": 8})
    assert run(ai.embed(["經驗"])) == run(ai.embed(["經驗"]))
    assert len(run(ai.embed([""]))[0]) == 8


def test_reflection_draft_retains_question_and_evidence():
    result = run(AIService().respond({"message": "整理", "history": [{"role": "user", "content": "資工還是資管？"}], "context": [{"id": "e1"}]}))
    assert result["reflection"]["question"] == "資工還是資管？"
    assert result["reflection"]["evidence_refs"] == ["e1"]
    assert result["ui"][0]["confirmed"] is False


def test_extract_utf8_and_unsupported():
    assert extract_text("中文".encode(), "text/markdown") == "中文"
    assert extract_text(b"media", "video/mp4") == ""


def test_extract_docx_paragraphs_and_tables():
    from io import BytesIO
    from docx import Document
    doc = Document()
    doc.add_paragraph("Synthetic experience")
    doc.add_table(rows=1, cols=1).cell(0, 0).text = "Evidence"
    stream = BytesIO()
    doc.save(stream)
    text = extract_text(stream.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    assert "Synthetic experience" in text and "Evidence" in text


def test_extract_pdf_and_reject_docx_zip_bomb():
    from io import BytesIO
    from zipfile import ZipFile, ZIP_DEFLATED
    from pypdf import PdfWriter
    pdf = PdfWriter()
    pdf.add_blank_page(width=100, height=100)
    stream = BytesIO()
    pdf.write(stream)
    assert extract_text(stream.getvalue(), "application/pdf") == ""
    archive = BytesIO()
    with ZipFile(archive, "w", ZIP_DEFLATED) as zip_file:
        zip_file.writestr("word/document.xml", "a" * 100000)
    with pytest.raises(ValueError, match="expansion"):
        extract_text(archive.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")


def test_counterevidence_is_preserved():
    result = run(AIService().respond({"message": "探索", "context": [{"id": "support"}, {"id": "counter", "counter_evidence": True}]}))
    assert result["hypothesis"]["evidence_refs"] == ["support"]
    assert result["hypothesis"]["counter_evidence_refs"] == ["counter"]


def test_provider_cannot_forge_confirmation_or_ui_hypothesis():
    bad = output(ui=[{"schema_version": "1.0", "type": "hypothesis_card", "hypothesis_id": "forged", "statement": "already confirmed", "evidence_refs": [], "uncertainty": "none"}])
    result = run(service(lambda request: completion(bad)).respond({"message": "探索"}))
    assert result["hypothesis"] is None and result["metadata"]["fallback"]
