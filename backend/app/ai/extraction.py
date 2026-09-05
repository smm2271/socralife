"""Bounded text extraction and page rendering for vision analysis."""
import base64
from io import BytesIO
from zipfile import ZipFile, BadZipFile

MAX_BYTES = 50 * 1024 * 1024
MAX_TEXT = 1_000_000


def extract_text(content: bytes, mime_type: str) -> str:
    if len(content) > MAX_BYTES:
        raise ValueError("File exceeds extraction limit")
    mime = mime_type.split(";", 1)[0].strip().lower()
    if mime in ("text/plain", "text/markdown"):
        return content.decode("utf-8-sig", errors="replace")[:MAX_TEXT].replace("\x00", "")
    if mime == "application/pdf":
        from pypdf import PdfReader
        reader = PdfReader(BytesIO(content), strict=True)
        if reader.is_encrypted:
            raise ValueError("Encrypted PDF text extraction is unavailable")
        if len(reader.pages) > 500:
            raise ValueError("PDF page limit exceeded")
        chunks, size = [], 0
        for page in reader.pages:
            stream = page.get_contents()
            if stream is not None and len(stream.get_data()) > 5_000_000:
                raise ValueError("PDF page content exceeds parsing limit")
            text = page.extract_text() or ""
            chunks.append(text[:MAX_TEXT-size])
            size += len(chunks[-1])
            if size >= MAX_TEXT:
                break
        return "\n".join(chunks)[:MAX_TEXT]
    if mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        from docx import Document
        try:
            with ZipFile(BytesIO(content)) as archive:
                entries = archive.infolist()
                if len(entries) > 5000 or sum(x.file_size for x in entries) > MAX_BYTES or any(x.file_size > max(x.compress_size, 1) * 200 for x in entries):
                    raise ValueError("DOCX expansion limit exceeded")
        except BadZipFile as exc:
            raise ValueError("Invalid DOCX container") from exc
        document = Document(BytesIO(content))
        lines = [p.text for p in document.paragraphs]
        lines.extend(cell.text for table in document.tables for row in table.rows for cell in row.cells)
        return "\n".join(lines)[:MAX_TEXT]
    return ""

def visual_pages(content: bytes, mime_type: str, max_pages: int = 8) -> list[dict]:
    """Render image/PDF pages as bounded data URIs for a vision provider."""
    if len(content) > MAX_BYTES: raise ValueError("File exceeds visual limit")
    mime = mime_type.split(";", 1)[0].strip().lower()
    pages: list[bytes] = []
    if mime in ("image/png", "image/jpeg"):
        pages = [content]
    elif mime == "application/pdf":
        import fitz
        document = fitz.open(stream=content, filetype="pdf")
        if document.is_encrypted or document.page_count > 500: raise ValueError("PDF visual limits exceeded")
        for page in list(document)[:max_pages]:
            pixmap = page.get_pixmap(matrix=fitz.Matrix(1.4, 1.4), alpha=False)
            pages.append(pixmap.tobytes("png"))
    else:
        return []
    return [{"data_uri": "data:image/png;base64," + base64.b64encode(blob).decode("ascii")} for blob in pages[:max_pages]]
